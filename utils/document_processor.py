"""
Document processing utilities for multiple file formats
"""
import os
import json
import tempfile
from typing import Dict, List, Tuple
import pandas as pd
from pypdf import PdfReader
from docx import Document
import streamlit as st

class DocumentProcessor:
    """Process various document formats for RAG pipeline"""
    
    @staticmethod
    def process_pdf(file) -> Tuple[str, Dict]:
        """Extract text from PDF file"""
        try:
            pdf_reader = PdfReader(file)
            text_content = []
            metadata = {
                'pages': len(pdf_reader.pages),
                'format': 'PDF'
            }
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                text_content.append(f"[Page {page_num}]\n{text}")
            
            full_text = "\n\n".join(text_content)
            return full_text, metadata
            
        except Exception as e:
            raise Exception(f"Error processing PDF: {str(e)}")
    
    @staticmethod
    def process_txt(file) -> Tuple[str, Dict]:
        """Extract text from TXT file"""
        try:
            text_content = file.read().decode('utf-8')
            metadata = {
                'format': 'TXT',
                'size': len(text_content)
            }
            return text_content, metadata
            
        except Exception as e:
            raise Exception(f"Error processing TXT: {str(e)}")
    
    @staticmethod
    def process_docx(file) -> Tuple[str, Dict]:
        """Extract text from DOCX file"""
        try:
            doc = Document(file)
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells])
                    text_content.append(row_text)
            
            full_text = "\n\n".join(text_content)
            metadata = {
                'format': 'DOCX',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables)
            }
            
            return full_text, metadata
            
        except Exception as e:
            raise Exception(f"Error processing DOCX: {str(e)}")
    
    @staticmethod
    def process_csv(file) -> Tuple[str, Dict]:
        """Extract text from CSV file"""
        try:
            df = pd.read_csv(file)
            
            # Convert DataFrame to text representation
            text_content = f"CSV Data Summary:\n"
            text_content += f"Columns: {', '.join(df.columns.tolist())}\n"
            text_content += f"Rows: {len(df)}\n\n"
            text_content += "Data:\n"
            text_content += df.to_string()
            
            metadata = {
                'format': 'CSV',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist()
            }
            
            return text_content, metadata
            
        except Exception as e:
            raise Exception(f"Error processing CSV: {str(e)}")
    
    @staticmethod
    def process_json(file) -> Tuple[str, Dict]:
        """Extract text from JSON file"""
        try:
            json_data = json.load(file)
            
            # Convert JSON to readable text
            text_content = json.dumps(json_data, indent=2)
            
            metadata = {
                'format': 'JSON',
                'size': len(text_content)
            }
            
            return text_content, metadata
            
        except Exception as e:
            raise Exception(f"Error processing JSON: {str(e)}")
    
    @classmethod
    def process_file(cls, uploaded_file) -> Tuple[str, Dict]:
        """
        Process uploaded file based on its type
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            
        Returns:
            Tuple of (text_content, metadata)
        """
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        processors = {
            'pdf': cls.process_pdf,
            'txt': cls.process_txt,
            'docx': cls.process_docx,
            'csv': cls.process_csv,
            'json': cls.process_json
        }
        
        if file_extension not in processors:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        processor = processors[file_extension]
        text_content, metadata = processor(uploaded_file)
        
        # Add common metadata
        metadata.update({
            'filename': uploaded_file.name,
            'file_size': uploaded_file.size,
            'file_type': file_extension
        })
        
        return text_content, metadata
    
    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for better context
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start = end - overlap
        
        return chunks
    
    @staticmethod
    def validate_file(uploaded_file, max_size_mb: int = 10) -> Tuple[bool, str]:
        """
        Validate uploaded file
        
        Args:
            uploaded_file: Streamlit UploadedFile object
            max_size_mb: Maximum file size in MB
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        max_size_bytes = max_size_mb * 1024 * 1024
        if uploaded_file.size > max_size_bytes:
            return False, f"File size exceeds {max_size_mb}MB limit"
        
        # Check file extension
        file_extension = uploaded_file.name.split('.')[-1].lower()
        supported_types = ['pdf', 'txt', 'docx', 'csv', 'json']
        
        if file_extension not in supported_types:
            return False, f"Unsupported file type. Supported: {', '.join(supported_types)}"
        
        return True, ""